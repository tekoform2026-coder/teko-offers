import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    """Задава цвят на фона на клетка от таблица."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_word_offer(client_name, project_name, elements, price_per_m2=15.00, offer_type="rental", filename="TEKO_Offer.docx"):
    """
    Генерира професионална търговска оферта в Word (.docx) за ПЛАСПАНЕЛ ООД / TEKO.
    offer_type: "rental" / "наем" или "sale" / "продажба"
    """
    doc = Document()

    # Полета на страницата
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    GREEN_COLOR = RGBColor(46, 125, 50)    # TEKO Зелено (лого)
    BLUE_COLOR = RGBColor(0, 51, 102)      # Корпоративно синьо
    BLUE_HEX = "003366"                     # Хекс цвят за шапката на таблицата

    is_sale = (str(offer_type).lower() in ["sale", "продажба", "закупуване"])

    # 1. Шапка - Фирмени данни вляво, Зелено лого TEKO вдясно
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.rows[0].cells[0].width = Inches(4.5)
    header_table.rows[0].cells[1].width = Inches(2.3)

    # Фирмени данни вляво
    cell_left = header_table.rows[0].cells[0]
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r_comp = p_left.add_run("ПЛАСПАНЕЛ ООД | ЕИК 208141542\n")
    r_comp.bold = True
    r_comp.font.size = Pt(10.5)
    r_comp.font.color.rgb = BLUE_COLOR

    r_addr = p_left.add_run(
        "2700 Благоевград, ул. „Ал. Стамболийски” №9, ет. 1\n"
        "www.tekoform.com | e-mail: bulgaria@tekoform.com | тел: +359 879 044 188"
    )
    r_addr.font.size = Pt(8.5)
    r_addr.font.color.rgb = RGBColor(90, 90, 90)

    # Лого TEKO в зелено вдясно
    cell_right = header_table.rows[0].cells[1]
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_logo = p_right.add_run("TEKO")
    r_logo.bold = True
    r_logo.font.size = Pt(28)
    r_logo.font.color.rgb = GREEN_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2. Заглавие на офертата
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    action_text = "закупуване" if is_sale else "отдаване под наем"
    title_run = title_p.add_run(f"ОФЕРТА\nЗа {action_text} на пластмасова кофражна система TEKO")
    title_run.bold = True
    title_run.font.size = Pt(15)
    title_run.font.color.rgb = BLUE_COLOR
    title_p.paragraph_format.space_after = Pt(10)

    # 3. Детайли за клиента и обекта
    details_p = doc.add_paragraph()
    details_p.paragraph_format.space_after = Pt(12)
    
    details_p.add_run("До: ").bold = True
    details_p.add_run(f"{client_name if client_name else '—'}\n")
    details_p.add_run("Относно: ").bold = True
    details_p.add_run(f"Кофриране на стоманобетонови елементи за обект: „{project_name if project_name else '—'}“\n")
    details_p.add_run("Дата: ").bold = True
    details_p.add_run(f"{datetime.now().strftime('%d.%m.%Y')} г.")

    # 4. Таблица с елементи и цени
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    
    col_last_title = "Обща сума (€)" if is_sale else "Наем (€)"
    headers = ['Елемент', 'Размери', 'Брой', 'Площ (m²)', 'Ед. цена (€/m²)', col_last_title]
    
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        set_cell_background(hdr_cells[i], BLUE_HEX)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)

    total_area = 0.0

    for elem in elements:
        row_cells = table.add_row().cells
        
        # Поддръжка на ключове на английски и български език
        elem_type = elem.get('type', elem.get('Елемент', ''))
        elem_dims = elem.get('dimensions', elem.get('Размери', ''))
        elem_count = elem.get('count', elem.get('Брой', 0))
        elem_area = elem.get('area_m2', elem.get('Площ (m²)', 0.0))

        row_cells[0].text = str(elem_type)
        row_cells[1].text = str(elem_dims)
        row_cells[2].text = f"{int(elem_count)} бр."
        row_cells[3].text = f"{float(elem_area):.2f} m²"
        row_cells[4].text = f"{price_per_m2:.2f} €"
        
        elem_cost = float(elem_area) * price_per_m2
        row_cells[5].text = f"{elem_cost:.2f} €"

        total_area += float(elem_area)

        for i in range(6):
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9)

    # 5. Финансово резюме
    subtotal = total_area * price_per_m2
    vat = subtotal * 0.20
    grand_total = subtotal + vat

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.line_spacing = 1.25
    summary_p.paragraph_format.space_after = Pt(12)
    
    summary_p.add_run("Обща кофражна площ: ").bold = True
    summary_p.add_run(f"{total_area:.2f} m²\n")
    
    action_label = "покупка" if is_sale else "наем"
    summary_p.add_run(f"Обща стойност за {action_label} (без ДДС): ").bold = True
    summary_p.add_run(f"{subtotal:.2f} €\n")
    summary_p.add_run("ДДС (20%): ").bold = True
    summary_p.add_run(f"{vat:.2f} €\n")
    
    grand_label = f"ОБЩО ЗА ПЛАЩАНЕ: {grand_total:.2f} €" if is_sale else f"ОБЩО ЗА ПЛАЩАНЕ / МЕСЕЦ: {grand_total:.2f} €"
    total_run = summary_p.add_run(grand_label)
    total_run.bold = True
    total_run.font.size = Pt(12)
    total_run.font.color.rgb = BLUE_COLOR

    # 6. Условия на офертата
    notes_p = doc.add_paragraph()
    notes_p.paragraph_format.space_before = Pt(10)
    
    run_th = notes_p.add_run("Условия на офертата:\n")
    run_th.bold = True
    run_th.font.color.rgb = BLUE_COLOR

    terms = [
        "1. Всички цени са посочени в евро (€) без включен ДДС.",
        "2. Начин на плащане: 100% авансово плащане при потвърждение на поръчката." if is_sale else "2. Минимален срок за наем: 30 календарни дни.",
        "3. Срок за доставка: До 5 работни дни след постъпване на плащането.",
        "4. Гаранция: 12 месеца за фабрични дефекти при спазване на инструкциите за работа.",
        "5. Забележка: В цената не са включени анкери, тапи, кофражно масло и пластмасови тръби.",
        "6. Място на вземане: Склад на фирмата (Транспортът е за сметка на Купувача/Наемателя)."
    ]
    for term in terms:
        r = notes_p.add_run(f"{term}\n")
        r.font.size = Pt(8.5)

    # 7. Подпис
    sign_p = doc.add_paragraph()
    sign_p.paragraph_format.space_before = Pt(12)
    sign_p.add_run("С уважение,\n").bold = True
    run_sign = sign_p.add_run("Екипът на ПЛАСПАНЕЛ ООД\nwww.tekoform.com")
    run_sign.font.color.rgb = BLUE_COLOR
    run_sign.bold = True

    doc.save(filename)
    return filename
