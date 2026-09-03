import io
import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
import streamlit as st
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from gemini_agent2 import analyze_blueprint
from drawing_generator import generate_wall_2d, generate_wall_3d, generate_pdf_drawings

st.set_page_config(
    page_title="TEKO - Вертикален Кофраж и Оферти",
    page_icon="🏗️",
    layout="wide"
)

def set_cell_background(cell, hex_color):
    """Задава цвят на фона на клетка от таблица."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def calculate_height_breakdown(height_cm):
    height_levels = []
    rem = height_cm
    while rem >= 150:
        height_levels.append(150)
        rem -= 150
    while rem >= 120:
        height_levels.append(120)
        rem -= 120
    while rem >= 60:
        height_levels.append(60)
        rem -= 60
    if rem > 0:
        height_levels.append(rem)
    return height_levels

def calculate_panel_width_breakdown(width_cm):
    panel_widths = [60, 35, 30, 25, 20]
    compensators = [15, 10, 5]
    
    remaining = width_cm
    result_panels = {}
    
    for w in panel_widths:
        count = remaining // w
        if count > 0:
            result_panels[f"TK _{int(w)}"] = int(count)
            remaining -= count * w
            
    for c in compensators:
        count = remaining // c
        if count > 0:
            result_panels[f"TC _{int(c)}"] = int(count)
            remaining -= count * c
            
    return result_panels

def get_element_teko_panels(elem_type, row):
    cnt = int(row.get("count", 1) or 1)
    h_m = float(row.get("height_m", 3.0) or 3.0)
    h_cm = h_m * 100
    h_levels = calculate_height_breakdown(h_cm)
    
    element_panels = {}
    
    def add_face_panels(face_width_cm):
        p_breakdown = calculate_panel_width_breakdown(face_width_cm)
        for h_val in h_levels:
            for p_code, p_cnt in p_breakdown.items():
                panel_name = p_code.replace("_", f"{int(h_val)}/")
                element_panels[panel_name] = element_panels.get(panel_name, 0) + p_cnt * 2 * cnt

    if elem_type == "column":
        w_cm = float(row.get("width_m", 0.3) or 0.3) * 100
        l_cm = float(row.get("length_m", 0.5) or 0.5) * 100
        add_face_panels(w_cm)
        add_face_panels(l_cm)
    elif elem_type == "wall":
        l_cm = float(row.get("length_m", 5.0) or 5.0) * 100
        add_face_panels(l_cm)
    elif elem_type == "l_wall":
        l1_cm = float(row.get("l1_m", 2.0) or 2.0) * 100
        l2_cm = float(row.get("l2_m", 2.0) or 2.0) * 100
        add_face_panels(l1_cm)
        add_face_panels(l2_cm)
    elif elem_type == "u_wall":
        l1_cm = float(row.get("l1_m", 2.0) or 2.0) * 100
        l2_cm = float(row.get("l2_m", 2.0) or 2.0) * 100
        l3_cm = float(row.get("l3_m", 2.0) or 2.0) * 100
        add_face_panels(l1_cm)
        add_face_panels(l2_cm)
        add_face_panels(l3_cm)
        
    return element_panels

def format_element_label(elem_type, name):
    name_str = str(name).strip() if name else ""
    if elem_type == "column":
        return name_str if name_str.lower().startswith("колона") else f"Колона {name_str}".strip()
    elif elem_type == "wall":
        return name_str if name_str.lower().startswith("стена") else f"Стена {name_str}".strip()
    elif elem_type == "l_wall":
        return name_str if (name_str.lower().startswith("l-стена") or name_str.lower().startswith("стена")) else f"L-Стена {name_str}".strip()
    elif elem_type == "u_wall":
        return name_str if (name_str.lower().startswith("u-ядро") or name_str.lower().startswith("ядро")) else f"U-Ядро {name_str}".strip()
    return name_str

def process_uploaded_file(uploaded_file):
    uploaded_file.seek(0)
    if uploaded_file.type == "application/pdf":
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        page = doc.load_page(0)
        pix = page.get_pixmap(dpi=150)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return img
    else:
        return Image.open(uploaded_file)

def generate_word_offer(client_name, project_name, offer_date, offer_type, price_per_m2, detailed_rows, total_area, subtotal, vat_amount, grand_total):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    GREEN_COLOR = RGBColor(46, 125, 50)
    BLUE_COLOR = RGBColor(0, 51, 102)
    BLUE_HEX = "003366"

    is_sale = ("закуп" in offer_type.lower() or "продажба" in offer_type.lower())

    p_logo = doc.add_paragraph()
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(2)
    r_logo = p_logo.add_run("TEKO")
    r_logo.bold = True
    r_logo.font.size = Pt(26)
    r_logo.font.color.rgb = GREEN_COLOR

    p_comp = doc.add_paragraph()
    p_comp.paragraph_format.space_after = Pt(12)
    p_comp.paragraph_format.line_spacing = 1.15

    r_comp = p_comp.add_run("ПЛАСПАНЕЛ ООД | ЕИК 208141542\n")
    r_comp.bold = True
    r_comp.font.size = Pt(10.5)
    r_comp.font.color.rgb = BLUE_COLOR

    r_addr = p_comp.add_run(
        "2700 Благоевград, ул. „Ал. Стамболийски” №9, ет. 1\n"
        "www.tekoform.com | e-mail: bulgaria@tekoform.com | тел: +359 879 044 188"
    )
    r_addr.font.size = Pt(8.5)
    r_addr.font.color.rgb = RGBColor(90, 90, 90)

    p_title = doc.add_paragraph()
    action_text = "закупуване" if is_sale else "наемане"
    run_title = p_title.add_run(f"ОФЕРТА\nЗа {action_text} на пластмасова кофражна система TEKO")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = BLUE_COLOR
    p_title.paragraph_format.space_after = Pt(10)

    p_info = doc.add_paragraph()
    p_info.paragraph_format.line_spacing = 1.25
    p_info.paragraph_format.space_after = Pt(12)
    p_info.add_run("До: ").bold = True
    p_info.add_run(f"{client_name}\n")
    p_info.add_run("Относно: ").bold = True
    p_info.add_run(f"Кофриране на стоманобетонови елементи за обект: „{project_name}“\n")
    p_info.add_run("Дата: ").bold = True
    p_info.add_run(f"{offer_date.strftime('%d.%m.%Y')} г.")

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    headers = ["Елемент", "Размери", "Брой", "Площ (m²)", "Ед. цена (€/m²)", "Обща сума (€)"]
    
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], BLUE_HEX)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)

    for row_data in detailed_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row_data["Елемент"])
        row_cells[1].text = str(row_data["Размери"])
        row_cells[2].text = f"{row_data['Брой']} бр."
        row_cells[3].text = f"{row_data['Площ (m²)']:.2f} m²"
        row_cells[4].text = f"{price_per_m2:.2f} €"
        row_cells[5].text = f"{row_data['Обща сума (€)']:.2f} €"
        
        for i, c in enumerate(row_cells):
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    action_word = "покупка" if is_sale else "наем"
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.line_spacing = 1.25
    p_summary.paragraph_format.space_after = Pt(12)
    
    p_summary.add_run(f"Обща кофражна площ: {total_area:.2f} m²\n").bold = True
    p_summary.add_run(f"Обща стойност за {action_word} (без ДДС): {subtotal:.2f} €\n").bold = True
    p_summary.add_run(f"ДДС (20%): {vat_amount:.2f} €\n")
    
    run_total = p_summary.add_run(f"ОБЩО ЗА ПЛАЩАНЕ: {grand_total:.2f} €")
    run_total.bold = True
    run_total.font.size = Pt(11)
    run_total.font.color.rgb = BLUE_COLOR

    notes_p = doc.add_paragraph()
    notes_p.paragraph_format.space_before = Pt(8)
    notes_p.paragraph_format.space_after = Pt(12)
    notes_p.paragraph_format.line_spacing = 1.15
    
    run_th = notes_p.add_run("Условия на офертата:\n")
    run_th.bold = True
    run_th.font.color.rgb = BLUE_COLOR

    terms = [
        "1. Всички цени са посочени в евро (€) без включен ДДС.",
        "2. Начин на плащане: 100% авансово плащане при потвърждение на поръчката.",
        "3. Срок за доставка: До 5 работни дни след постъпване на плащането.",
        "4. Гаранция: 12 месеца за фабрични дефекти при спазване на инструкциите за работа.",
        "5. Забележка: В цената не са включени анкери, тапи, кофражно масло и пластмасови тръби.",
        "6. Място на вземане: Склад на фирмата (Транспортът е за сметка на Купувача/Наемателя)."
    ]
    for term in terms:
        r = notes_p.add_run(f"{term}\n")
        r.font.size = Pt(8.5)

    p_sign = doc.add_paragraph()
    p_sign.paragraph_format.space_before = Pt(8)
    p_sign.add_run("С уважение,\n").bold = True
    run_sign = p_sign.add_run("Екипът на ПЛАСПАНЕЛ ООД\nwww.tekoform.com")
    run_sign.font.color.rgb = BLUE_COLOR
    run_sign.bold = True

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

api_key = st.secrets.get("GEMINI_API_KEY", "")

if "blueprint_data" not in st.session_state:
    st.session_state["blueprint_data"] = None
if "used_model" not in st.session_state:
    st.session_state["used_model"] = None
if "edited_df" not in st.session_state:
    st.session_state["edited_df"] = pd.DataFrame()
if "render_3d" not in st.session_state:
    st.session_state["render_3d"] = False

with st.sidebar:
    st.header("⚙️ Настройки")
    if not api_key:
        api_key = st.text_input("Въведете Gemini API Key:", type="password")
    else:
        st.success("🔑 API Ключът е зареден!")

st.title("🏗️ ПЛАСПАНЕЛ ООД - Пластмасова Кофражна Система TEKO")

tab1, tab2, tab3, tab4 = st.tabs([
    "📐 Чертеж и Редактиране", 
    "🧱 Кофражни Елементи и Панели TEKO", 
    "🎨 2D/3D Чертежи",
    "📄 Оферта"
])

with tab1:
    st.header("1. Качване на чертеж и AI разчитане")
    uploaded_file = st.file_uploader("Изберете чертеж (PDF, PNG, JPG):", type=["pdf", "png", "jpg", "jpeg"])

    if uploaded_file:
        col_img, col_actions = st.columns([1, 1])
        with col_img:
            st.subheader("🖼️ Преглед на файла")
            try:
                processed_img = process_uploaded_file(uploaded_file)
                st.image(processed_img, use_container_width=True, caption="Качен чертеж")
            except Exception as e:
                st.error(f"Грешка при зареждането: {e}")
                processed_img = None

        with col_actions:
            st.subheader("🤖 AI Разчитане")
            if not api_key:
                st.warning("⚠️ Моля, въведете Gemini API Key в страничното меню.")
            else:
                if st.button("🔍 Разчети чертежа с Vision AI", type="primary", use_container_width=True):
                    if processed_img:
                        with st.spinner("Извличане на вертикални кофражни елементи..."):
                            try:
                                res = analyze_blueprint(processed_img, api_key)
                                if isinstance(res, tuple):
                                    blueprint_data, used_model = res
                                else:
                                    blueprint_data, used_model = res, "Gemini Vision"

                                st.session_state["blueprint_data"] = blueprint_data
                                st.session_state["used_model"] = used_model
                                elements_list = blueprint_data.get("elements", [])
                                st.session_state["edited_df"] = pd.DataFrame(elements_list)
                                st.success(f"✅ Готово! Разчетено с: **{used_model}**")
                            except Exception as e:
                                st.error(f"Грешка при анализа: {e}")

    if not st.session_state["edited_df"].empty:
        st.divider()
        st.subheader("✏️ Корекция на разчетените вертикални елементи")
        edited_df = st.data_editor(
            st.session_state["edited_df"],
            num_rows="dynamic",
            column_config={
                "type": st.column_config.SelectboxColumn("Тип елемент", options=["column", "wall", "l_wall", "u_wall"], required=True),
                "name": st.column_config.TextColumn("Наименование / Маркировка"),
                "count": st.column_config.NumberColumn("Брой", min_value=1, step=1, default=1),
                "width_m": st.column_config.NumberColumn("Ширина (м)", format="%.2f"),
                "length_m": st.column_config.NumberColumn("Дължина (м)", format="%.2f"),
                "thickness_m": st.column_config.NumberColumn("Дебелина (м)", format="%.2f"),
                "l1_m": st.column_config.NumberColumn("Рамо 1 (м)", format="%.2f"),
                "l2_m": st.column_config.NumberColumn("Рамо 2 (м)", format="%.2f"),
                "l3_m": st.column_config.NumberColumn("Рамо 3 (м)", format="%.2f"),
                "height_m": st.column_config.NumberColumn("Височина (м)", format="%.2f")
            },
            use_container_width=True,
            key="elements_editor_tab1"
        )
        st.session_state["edited_df"] = edited_df

with tab2:
    st.header("2. Спецификация на кофражните елементи и панели TEKO")
    df_calc = st.session_state["edited_df"]
    
    if df_calc.empty:
        st.info("ℹ️ Качете чертеж в Таб 1 или въведете елементи ръчно.")
    else:
        detailed_rows = []
        total_a = 0.0
        project_panels_summary = {}

        for _, row in df_calc.iterrows():
            cnt = int(row.get("count", 1) or 1)
            h = float(row.get("height_m", 3.0) or 3.0)
            elem_type = str(row.get("type", "wall"))
            name = str(row.get("name", "Елемент"))

            area = 0.0
            dim_str = ""
            elem_label = format_element_label(elem_type, name)

            if elem_type == "column":
                w = float(row.get("width_m", 0.3) or 0.3)
                l = float(row.get("length_m", 0.5) or 0.5)
                area = 2 * (w + l) * h * cnt
                dim_str = f"{int(w*100)}x{int(l*100)} cm, H={h:.1f}m"

            elif elem_type == "wall":
                l = float(row.get("length_m", 5.0) or 5.0)
                t = float(row.get("thickness_m", 0.25) or 0.25)
                area = 2 * l * h * cnt
                dim_str = f"L={l:.1f}m, B={int(t*100)}cm, H={h:.1f}m"

            elif elem_type == "l_wall":
                l1 = float(row.get("l1_m", 2.0) or 2.0)
                l2 = float(row.get("l2_m", 2.0) or 2.0)
                t = float(row.get("thickness_m", 0.25) or 0.25)
                area = 2 * (l1 + l2) * h * cnt
                dim_str = f"L={l1:.1f}+{l2:.1f}m, B={int(t*100)}cm, H={h:.1f}m"

            elif elem_type == "u_wall":
                l1 = float(row.get("l1_m", 2.0) or 2.0)
                l2 = float(row.get("l2_m", 2.0) or 2.0)
                l3 = float(row.get("l3_m", 2.0) or 2.0)
                t = float(row.get("thickness_m", 0.25) or 0.25)
                area = 2 * (l1 + l2 + l3) * h * cnt
                dim_str = f"L={l1:.1f}+{l2:.1f}+{l3:.1f}m, B={int(t*100)}cm, H={h:.1f}m"

            total_a += area
            panels_dict = get_element_teko_panels(elem_type, row)
            for p_name, p_qty in panels_dict.items():
                project_panels_summary[p_name] = project_panels_summary.get(p_name, 0) + p_qty

            panels_str = ", ".join([f"{k} ({v} бр.)" for k, v in panels_dict.items()])

            detailed_rows.append({
                "Елемент": elem_label,
                "Размери": dim_str,
                "Брой": cnt,
                "Площ (m²)": round(area, 2),
                "Панели TEKO (Вид и брой)": panels_str
            })

        df_detailed = pd.DataFrame(detailed_rows)
        st.metric("📊 ОБЩА КОФРАЖНА ПЛОЩ", f"{total_a:.2f} m²")
        st.subheader("📋 Спецификация на кофражните елементи и съответните панели")
        st.dataframe(df_detailed, use_container_width=True)

        st.divider()
        st.subheader("📦 Общ брой нужни панели TEKO за целия обект")
        df_panels_sum = pd.DataFrame([
            {"Код на панела / коф. елемент": k, "Общ брой (бр.)": v} 
            for k, v in sorted(project_panels_summary.items())
        ])
        st.dataframe(df_panels_sum, use_container_width=True)

with tab3:
    st.header("3. 2D Развертки и 3D Моделиране")
    df_calc = st.session_state["edited_df"]

    if df_calc.empty:
        st.info("ℹ️ Няма въведени елементи. Качете чертеж в Таб 1 или въведете данни ръчно.")
    else:
        selected_idx = st.selectbox(
            "Изберете елемент за визуализация:",
            options=list(range(len(df_calc))),
            format_func=lambda i: f"{format_element_label(df_calc.iloc[i].get('type', 'wall'), df_calc.iloc[i].get('name', 'Елемент'))}"
        )

        if st.button("🏗️ Генерирай 3D модел и PDF чертежи", type="primary"):
            st.session_state["render_3d"] = True

        if st.session_state["render_3d"]:
            row = df_calc.iloc[selected_idx]
            elem_type = str(row.get("type", "wall"))
            name = str(row.get("name", "Елемент"))
            h_cm = float(row.get("height_m", 3.0) or 3.0) * 100
            t_cm = float(row.get("thickness_m", 0.25) or 0.25) * 100

            col_2d, col_3d = st.columns(2)

            with col_2d:
                st.subheader("📐 2D Развертка на панелите")
                l_cm = float(row.get("length_m", 5.0) or 5.0) * 100
                if elem_type in ["l_wall", "u_wall"]:
                    l_cm = float(row.get("l1_m", 2.0) or 2.0) * 100
                fig_2d = generate_wall_2d(l_cm, h_cm, wall_name=format_element_label(elem_type, name))
                st.image(fig_2d, use_container_width=True)

            with col_3d:
                st.subheader("🧊 3D Модел")
                l1_cm = float(row.get("l1_m", 2.0) or 2.0) * 100 if "l1_m" in row and pd.notnull(row["l1_m"]) else float(row.get("length_m", 5.0) or 5.0) * 100
                l2_cm = float(row.get("l2_m", 2.0) or 2.0) * 100 if "l2_m" in row and pd.notnull(row["l2_m"]) else 150
                l3_cm = float(row.get("l3_m", 2.0) or 2.0) * 100 if "l3_m" in row and pd.notnull(row["l3_m"]) else 150

                wall_type_map = {
                    "wall": "Права стена",
                    "column": "Права стена",
                    "l_wall": "L-образна стена",
                    "u_wall": "U-образна стена"
                }

                fig_3d = generate_wall_3d(
                    wall_type=wall_type_map.get(elem_type, "Права стена"),
                    dim_a=l1_cm,
                    dim_b=l2_cm,
                    dim_c=l3_cm,
                    height=h_cm,
                    thickness=t_cm
                )
                st.plotly_chart(fig_3d, use_container_width=True)

            st.divider()
            st.subheader("📄 Генериране на PDF с чертежи")

            pdf_elements = []
            bom_summary = {}
            for _, r in df_calc.iterrows():
                e_type = str(r.get("type", "wall"))
                e_name = format_element_label(e_type, str(r.get("name", "Елемент")))
                e_h = float(r.get("height_m", 3.0) or 3.0) * 100
                e_l = float(r.get("length_m", 5.0) or 5.0) * 100
                e_t = float(r.get("thickness_m", 0.25) or 0.25) * 100
                if e_type in ["l_wall", "u_wall"]:
                    e_l = float(r.get("l1_m", 2.0) or 2.0) * 100
                
                pdf_elements.append({
                    "name": e_name, 
                    "type": e_type,
                    "length_a_cm": e_l, 
                    "height_cm": e_h,
                    "thickness_cm": e_t
                })

                p_dict = get_element_teko_panels(e_type, r)
                for pk, pv in p_dict.items():
                    bom_summary[pk] = bom_summary.get(pk, 0) + pv

            proj_info = {
                "client": st.session_state.get("client_name_in_tab", "Клиент"),
                "project": st.session_state.get("project_name_in_tab", "Обект TEKO")
            }

            pdf_bytes = generate_pdf_drawings(pdf_elements, bom_summary, proj_info)
            st.download_button(
                label="⬇️ Свали PDF чертежи и количествена сметка",
                data=pdf_bytes,
                file_name="Teko_Drawings.pdf",
                mime="application/pdf",
                type="primary"
            )

with tab4:
    st.header("Оферта")
    
    col_info1, col_info2 = st.columns(2)
    with col_info1:
        client_name = st.text_input("Име на клиента / Фирма:", value="", key="client_name_in_tab")
        project_name_input = st.text_input("Име на обект:", value="", key="project_name_in_tab")
    
    with col_info2:
        offer_type = st.selectbox("Тип на офертата:", ["Наем", "Продажба"], key="offer_type_in_tab")
        offer_date = st.date_input("Дата на офертата", key="offer_date_in_tab")

    st.markdown("### 💶 Настройка на цена")
    col_price1, col_price2 = st.columns(2)
    with col_price1:
        price_formwork = st.number_input("Цена за кофраж (€/m²):", min_value=0.0, value=15.0, step=0.5, format="%.2f", key="price_in_tab")
    with col_price2:
        vat_percent = st.number_input("ДДС (%):", min_value=0.0, value=20.0, step=1.0, format="%.1f", key="vat_in_tab")

    st.divider()

    df_calc = st.session_state["edited_df"]

    if df_calc.empty:
        st.info("ℹ️ Няма въведени елементи за изчисляване на оферта. Качете чертеж или въведете данни в Таб 1.")
    else:
        detailed_rows = []
        total_a = 0.0

        for _, row in df_calc.iterrows():
            cnt = int(row.get("count", 1) or 1)
            h = float(row.get("height_m", 3.0) or 3.0)
            elem_type = str(row.get("type", "wall"))
            name = str(row.get("name", "Елемент"))

            area = 0.0
            dim_str = ""
            elem_label = format_element_label(elem_type, name)

            if elem_type == "column":
                w = float(row.get("width_m", 0.3) or 0.3)
                l = float(row.get("length_m", 0.5) or 0.5)
                area = 2 * (w + l) * h * cnt
                dim_str = f"{int(w*100)}x{int(l*100)} cm, H={h:.1f}m"

            elif elem_type == "wall":
                l = float(row.get("length_m", 5.0) or 5.0)
                t = float(row.get("thickness_m", 0.25) or 0.25)
                area = 2 * l * h * cnt
                dim_str = f"L={l:.1f}m, B={int(t*100)}cm, H={h:.1f}m"

            elif elem_type == "l_wall":
                l1 = float(row.get("l1_m", 2.0) or 2.0)
                l2 = float(row.get("l2_m", 2.0) or 2.0)
                t = float(row.get("thickness_m", 0.25) or 0.25)
                area = 2 * (l1 + l2) * h * cnt
                dim_str = f"L={l1:.1f}+{l2:.1f}m, B={int(t*100)}cm, H={h:.1f}m"

            elif elem_type == "u_wall":
                l1 = float(row.get("l1_m", 2.0) or 2.0)
                l2 = float(row.get("l2_m", 2.0) or 2.0)
                l3 = float(row.get("l3_m", 2.0) or 2.0)
                t = float(row.get("thickness_m", 0.25) or 0.25)
                area = 2 * (l1 + l2 + l3) * h * cnt
                dim_str = f"L={l1:.1f}+{l2:.1f}+{l3:.1f}m, B={int(t*100)}cm, H={h:.1f}m"

            total_a += area
            item_cost = area * price_formwork

            detailed_rows.append({
                "Елемент": elem_label,
                "Размери": dim_str,
                "Брой": cnt,
                "Площ (m²)": area,
                "Ед. цена (€/m²)": price_formwork,
                "Обща сума (€)": item_cost
            })

        subtotal = total_a * price_formwork
        vat_amount = subtotal * (vat_percent / 100.0)
        grand_total = subtotal + vat_amount

        action_text = "наемане" if "наем" in offer_type.lower() else "закупуване"

        st.markdown("<h2 style='color: #2E7D32; margin-bottom: 0;'>TEKO</h2>", unsafe_allow_html=True)
        st.subheader("ПЛАСПАНЕЛ ООД | ЕИК 208141542")
        st.caption("2700 Благоевград, ул. „Ал. Стамболийски” №9, ет. 1 | www.tekoform.com | bulgaria@tekoform.com | тел: +359 879 044 188")
        st.markdown(f"### **ОФЕРТА**\n**За {action_text} на пластмасова кофражна система TEKO**")
        
        st.write(f"**До:** {client_name if client_name else '—'}")
        st.write(f"**Относно:** Кофриране на стоманобетонови елементи за обект: „{project_name_input if project_name_input else '—'}“")
        st.write(f"**Дата:** {offer_date.strftime('%d.%m.%Y')} г.")

        preview_df = pd.DataFrame(detailed_rows)
        preview_df["Площ (m²)"] = preview_df["Площ (m²)"].apply(lambda x: f"{x:.2f} m²")
        preview_df["Ед. цена (€/m²)"] = preview_df["Ед. цена (€/m²)"].apply(lambda x: f"{x:.2f} €")
        preview_df["Обща сума (€)"] = preview_df["Обща сума (€)"].apply(lambda x: f"{x:.2f} €")
        
        st.table(preview_df)

        action_name = "наем" if "наем" in offer_type.lower() else "покупка"

        st.markdown(f"**Обща кофражна площ:** {total_a:.2f} m²")
        st.markdown(f"**Обща стойност за {action_name} (без ДДС):** {subtotal:.2f} €")
        st.markdown(f"**ДДС ({vat_percent:.0f}%):** {vat_amount:.2f} €")
        st.markdown(f"### **ОБЩО ЗА ПЛАЩАНЕ: {grand_total:.2f} €**")

        st.divider()

        docx_file = generate_word_offer(
            client_name=client_name,
            project_name=project_name_input,
            offer_date=offer_date,
            offer_type=offer_type,
            price_per_m2=price_formwork,
            detailed_rows=detailed_rows,
            total_area=total_a,
            subtotal=subtotal,
            vat_amount=vat_amount,
            grand_total=grand_total
        )

        st.download_button(
            label="📄 Свали офертата във MS Word (.docx)",
            data=docx_file,
            file_name=f"Oferta_TEKO_{client_name.replace(' ', '_') if client_name else 'клиент'}_{offer_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
